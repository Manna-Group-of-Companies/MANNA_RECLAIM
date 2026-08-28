"""
The plant's paper log sheets, one page per line, generated from the app's own
field list.

    python build_logsheets.py                     -> Manna-Reclaim-Log-Sheets.docx
    python build_logsheets.py --logo manna.png    -> with the logo in the header

WHY THIS IS A SCRIPT AND NOT A DOCUMENT. The sheets have to ask for exactly what
the tablet asks for, or the paper and the record drift apart and the shift has to
be entered twice from two sources that disagree. Every column below is a field
the app actually stores - see server/src/validations/run.validation.js and the
per-kind sheets in client/src/pages/user/MachinesPage.tsx. When a field is added
to the app, it is added here, and the sheets are regenerated.

The one exception is the boiler, which is marked as such on its own page: there
is no boiler in the app at all. Its sheet is drawn from what a firewood-fired
boiler attendant's log ordinarily carries, and it is a first draft to be
corrected rather than a mirror of anything.

The output is a Word file on purpose - it is meant to be edited.
"""
import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# The plant's own colours, off the logo.
ORANGE = RGBColor(0xF1, 0x6A, 0x21)
CHARCOAL = RGBColor(0x44, 0x44, 0x44)
RULE = "999999"
HEAD_FILL = "EFEFEF"

BODY_FONT = "Calibri"

# A4 landscape, because the plant is in India and python-docx defaults to US
# Letter - which is 1.8 cm narrower and silently pushed the last column of every
# sheet off the page.
PAGE_W = Cm(29.7)
PAGE_H = Cm(21.0)
SIDE_MARGIN = Cm(1.0)
USABLE_CM = 29.7 - 2 * 1.0


# A landscape A4 page, less the top and bottom margins. Every sheet has to fit
# inside this: a log sheet that runs onto a second page is two sheets to file,
# and the second one is the half nobody signs.
PAGE_BUDGET_CM = 21.0 - 0.9 - 0.9
_used = [0.0]


def spend(cm):
    _used[0] += cm


def start_sheet():
    _used[0] = 0.0


def check_sheet(name):
    used = _used[0]
    room = PAGE_BUDGET_CM - used
    flag = "  <-- WILL SPILL ONTO A SECOND PAGE" if room < 0 else ""
    print(f"  {name:<12} {used:5.1f} cm of {PAGE_BUDGET_CM:.1f}{flag}")
    return room >= 0


def fit(widths):
    """
    Scale a list of column widths to exactly the printable width.

    Written once, here, rather than by adding up centimetres by hand at each
    table: a sheet whose columns total more than the page does not warn anybody,
    it just prints with the remarks column hanging off the right edge, and the
    first time anyone finds out is when a shift has been filled in on it.
    """
    total = sum(widths)
    return [w * USABLE_CM / total for w in widths]


def shade(cell, hex_fill):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def borders(table, sz=4, colour=RULE):
    """Every edge, because a log sheet is filled in by hand against the lines."""
    tbl = table._tbl
    pr = tbl.tblPr
    marks = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:color"), colour)
        marks.append(e)
    pr.append(marks)


def repeat_header(row):
    """A table that runs onto a second page keeps its column names."""
    pr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    pr.append(el)


def set_text(cell, text, *, size=8, bold=False, colour=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = BODY_FONT
    if colour is not None:
        run.font.color.rgb = colour
    return p


def row_height(row, cm):
    """Hand-filled rows need room for a pencil, not just for the font."""
    tr = row._tr
    pr = tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(int(cm * 567)))
    h.set(qn("w:hRule"), "atLeast")
    pr.append(h)


def landscape(doc, first=False):
    section = doc.sections[0] if first else doc.add_section()
    section.orientation = WD_ORIENT.LANDSCAPE
    # Set outright rather than swapped off the default, which is US Letter.
    section.page_width = PAGE_W
    section.page_height = PAGE_H
    section.top_margin = Cm(0.9)
    section.bottom_margin = Cm(0.9)
    section.left_margin = SIDE_MARGIN
    section.right_margin = SIDE_MARGIN
    return section


def masthead(doc, title, subtitle, logo):
    """
    The header every sheet carries: who, what, and the shift it belongs to.

    A three-column table rather than a Word header, because the shift boxes have
    to be written in and a header is awkward to fill on paper and worse to edit
    in Word.
    """
    spend(2.1)
    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    widths = tuple(Cm(w) for w in fit([6.4, 12.6, 8.5]))
    for i, w in enumerate(widths):
        t.columns[i].width = w
    cells = t.rows[0].cells
    for i, w in enumerate(widths):
        cells[i].width = w

    if logo and Path(logo).exists():
        cells[0].text = ""
        p = cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run().add_picture(str(logo), width=Cm(5.4))
    else:
        # A named, sized box rather than nothing: the sheet is printed and filed
        # whether or not the file was to hand, and an empty corner reads as a
        # design that forgot the logo.
        set_text(cells[0], "", size=9)
        for line in ("", "[  MANNA  GROUP  LOGO  ]", "paste the logo here, or re-run",
                     "the generator with --logo <file>", ""):
            para = cells[0].add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            r = para.add_run(line)
            r.font.size = Pt(9 if line.startswith("[") else 6.5)
            r.font.bold = line.startswith("[")
            r.font.name = BODY_FONT
            r.font.color.rgb = CHARCOAL

    cells[1].text = ""
    p = set_text(cells[1], "MANNA RUBBER PRODUCTS (P) LTD.", size=13, bold=True,
                 colour=CHARCOAL, align=WD_ALIGN_PARAGRAPH.CENTER)
    p2 = cells[1].add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run(title)
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.name = BODY_FONT
    r.font.color.rgb = ORANGE
    p3 = cells[1].add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(subtitle)
    r3.font.size = Pt(8)
    r3.font.name = BODY_FONT
    r3.font.color.rgb = CHARCOAL

    set_text(cells[2], "Date: ......................................", size=10)
    for line in ("Shift:    ☐ Day  (08:30 – 20:30)      ☐ Night (20:30 – 08:30)",
                 "Supervisor: ..........................................",
                 "Sheet ......... of .........."):
        p = cells[2].add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.size = Pt(9)
        run.font.name = BODY_FONT
    return t


def grid(doc, columns, rows, *, widths=None, head_size=7.5, row_cm=0.68):
    """One table: the column names, then that many empty lines to write on."""
    t = doc.add_table(rows=1 + rows, cols=len(columns))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    borders(t)
    if widths:
        widths = fit(widths)
        for i, w in enumerate(widths):
            t.columns[i].width = Cm(w)

    head = t.rows[0]
    repeat_header(head)
    row_height(head, 0.8)
    for i, name in enumerate(columns):
        c = head.cells[i]
        if widths:
            c.width = Cm(widths[i])
        shade(c, HEAD_FILL)
        set_text(c, name, size=head_size, bold=True, colour=CHARCOAL,
                 align=WD_ALIGN_PARAGRAPH.CENTER)

    spend(0.8 + rows * row_cm)
    for r in range(1, 1 + rows):
        row_height(t.rows[r], row_cm)
        if widths:
            for i, w in enumerate(widths):
                t.rows[r].cells[i].width = Cm(w)
    return t


def note(doc, text, *, size=7.5, italic=True, space_before=4):
    # Roughly one line per 105 characters at this size and column width.
    spend(0.34 * max(1, (len(text) // 105) + 1) + space_before / 28.35)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.italic = italic
    run.font.name = BODY_FONT
    run.font.color.rgb = CHARCOAL
    return p


def section_label(doc, text):
    spend(0.62)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.name = BODY_FONT
    run.font.color.rgb = ORANGE
    return p


def bearings(doc, machines, every_hours, kind_word):
    """
    The temperatures, on the machines that have them.

    Four positions per machine and a reading every two or three hours depending
    on the line, flagged over 80 °C - which is the app's own rule, so the paper
    asks for the same readings at the same cadence. PR1, R1, R2 and Grinder 1
    run on bushes rather than bearings and the sheet says so, because the crew
    calls them different things and a column headed for the wrong one gets left
    blank.
    """
    section_label(doc, f"{kind_word} temperatures (°C) — every {every_hours} hours, "
                       f"flag anything over 80 °C")
    cols = ["Time", "Machine", "Pos 1", "Pos 2", "Pos 3", "Pos 4", "Over 80 °C? Action taken"]
    widths = [2.2, 4.0, 2.0, 2.0, 2.0, 2.0, 13.3]
    grid(doc, cols, 5, widths=widths, row_cm=0.6)


def signoff(doc):
    section_label(doc, "Shift remarks")
    grid(doc, ["Time", "Stoppage / breakdown / anything the next shift needs to know"],
         2, widths=[2.6, 24.9], row_cm=0.62)

    spend(1.6)
    t = doc.add_table(rows=1, cols=3)
    t.autofit = False
    for i, w in enumerate(fit([9.2, 9.2, 9.1])):
        t.columns[i].width = Cm(w)
        t.rows[0].cells[i].width = Cm(w)
    labels = ("Operator", "Shift supervisor", "Checked by (Manager)")
    for i, label in enumerate(labels):
        c = t.rows[0].cells[i]
        set_text(c, "\n\n................................................", size=9)
        p = c.add_paragraph()
        run = p.add_run(label)
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.name = BODY_FONT
        run.font.color.rgb = CHARCOAL
    note(doc, "Entered in the tablet by: ..............................................   "
              "Date/time entered: ..............................   "
              "File this sheet with the shift it belongs to.", italic=False, size=8)


# ---------------------------------------------------------------------------
# The six sheets. Each column below is a field the app stores.
# ---------------------------------------------------------------------------

def special_line(doc, logo):
    masthead(doc, "SPECIAL LINE — SHIFT LOG",
             "Pre-Refiner 2 · Refiner 1 · Refiner 3 · Refiner 4 "
             "(and PR1 / Refiner 2 when they are turned onto the special line)", logo)
    note(doc, "One row per pass. Only the finishing pass is weighed - leave the weight blank "
              "on the others. Never write two batch numbers in one box.", space_before=3)

    cols = ["Batch No.", "Mixed with", "Formulation", "Grade", "Machine", "Pass no.",
            "Start time", "Stop time", "Crew", "Elec. meter start", "Elec. meter end",
            "Units (kWh)", "Hour meter start", "Hour meter end", "Hours run",
            "Weight out (kg)", "Remarks"]
    widths = [1.9, 1.7, 2.1, 1.7, 1.7, 1.0, 1.4, 1.4, 0.9, 1.9, 1.9, 1.4, 1.9, 1.9, 1.2, 1.8, 2.7]
    grid(doc, cols, 8, widths=widths)

    bearings(doc, ["PR2", "R1", "R3", "R4"], 3, "Bearing / bush")
    signoff(doc)


def coarse_line(doc, logo):
    masthead(doc, "COARSE LINE — SHIFT LOG",
             "Pre-Refiner 1 · Refiner 2 — worked as one line", logo)

    cols = ["Machine", "Start time", "Stop time", "Crew", "Elec. meter start", "Elec. meter end",
            "Units (kWh)", "Hour meter start", "Hour meter end", "Hours run",
            "Weight out (kg)", "Remarks"]
    widths = [2.6, 2.0, 2.0, 1.2, 2.5, 2.5, 1.9, 2.5, 2.5, 1.7, 2.3, 3.8]
    grid(doc, cols, 6, widths=widths)

    section_label(doc, "Charges fed to the line this shift")
    grid(doc, ["Vessel", "Formulation", "Capacity (kg)", "Discharged at", "Remarks"], 2,
         widths=[3.0, 5.0, 3.5, 3.5, 12.5], row_cm=0.62)

    bearings(doc, ["PR1", "R2"], 3, "Bush / bearing")
    signoff(doc)


def grinders(doc, logo):
    masthead(doc, "GRINDING LINE — SHIFT LOG",
             "Grinder 1 · Grinder 2 · Soorya Grinder", logo)
    note(doc, "Tyre fed: Truck (30#) or Bike (20#). The Soorya Grinder has no meters - leave "
              "those four columns blank on it.", space_before=3)

    cols = ["Machine", "Tyre fed", "Mesh", "Start time", "Stop time", "Crew",
            "Elec. meter start", "Elec. meter end", "Units (kWh)",
            "Hour meter start", "Hour meter end", "Hours run", "Weight out (kg)", "Remarks"]
    widths = [2.4, 1.9, 1.3, 1.8, 1.8, 1.0, 2.2, 2.2, 1.6, 2.2, 2.2, 1.5, 2.0, 3.4]
    grid(doc, cols, 8, widths=widths)

    bearings(doc, ["Grinder 1", "Grinder 2", "Soorya"], 2, "Bush / bearing")
    signoff(doc)


def cracker(doc, logo):
    masthead(doc, "CRACKER — SHIFT LOG", "Tyre cracking, and the picking gang that feeds it", logo)

    cols = ["Start time", "Stop time", "Crew", "Elec. meter start", "Elec. meter end",
            "Units (kWh)", "Hour meter start", "Hour meter end", "Hours run", "Remarks"]
    widths = [2.4, 2.4, 1.4, 3.0, 3.0, 2.2, 3.0, 3.0, 2.0, 5.1]
    grid(doc, cols, 5, widths=widths)

    section_label(doc, "Picking gang — pulling scrap tyres out of the yard")
    grid(doc, ["No. of labourers", "Hours (approx.)", "Why the gang was this size / what they did"],
         3, widths=[4.0, 4.0, 19.5], row_cm=0.62)

    bearings(doc, ["Cracker"], 2, "Bearing")
    signoff(doc)


def autoclaves(doc, logo):
    masthead(doc, "AUTOCLAVES — SHIFT LOG", "Autoclave A · Autoclave M — one row per charge", logo)

    cols = ["Vessel", "Batch No.", "Formulation", "Capacity (kg)", "Paired?",
            "Loaded date", "Loaded time", "Pressure reached at", "Door opened at",
            "Discharged at", "Firewood (kg)", "Crew", "Remarks"]
    widths = [1.8, 1.9, 3.0, 1.9, 1.4, 1.9, 1.7, 2.4, 2.1, 2.0, 1.9, 1.0, 4.5]
    grid(doc, cols, 7, widths=widths)

    note(doc, "Paired? — tick when both vessels were charged together and one crew attended both. "
              "The labour is shared, so it is counted as half on each.", space_before=1)

    section_label(doc, "Firewood")
    grid(doc, ["Firewood taken into the shift (kg)", "Firewood received (kg)",
               "Firewood left at end of shift (kg)", "Remarks"],
         1, widths=[6.5, 6.5, 6.5, 8.0], row_cm=0.62)
    signoff(doc)


def boiler(doc, logo):
    masthead(doc, "BOILER — SHIFT LOG", "Steam to the autoclaves — hourly readings", logo)

    spend(1.5)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(
        "NOTE — this is the one sheet that is not taken from the app. There is no boiler in "
        "MANNA RECLAIM: nothing on it is recorded anywhere at present, so the columns below are "
        "what a firewood-fired boiler log ordinarily carries rather than a copy of what the "
        "plant already does. Correct it to match how the boiler is actually attended, and it can "
        "be built into the app afterwards."
    )
    run.font.size = Pt(8.5)
    run.font.bold = True
    run.font.name = BODY_FONT
    run.font.color.rgb = ORANGE

    cols = ["Time", "Steam pressure (kg/cm²)", "Water level", "Feed pump (on/off)",
            "Firewood added (kg)", "Flue gas / furnace", "Blowdown done", "Attendant", "Remarks"]
    widths = [2.0, 3.4, 2.6, 2.8, 3.0, 3.0, 2.6, 2.6, 5.5]
    grid(doc, cols, 10, widths=widths, row_cm=0.62)

    section_label(doc, "Start and end of shift checks")
    grid(doc, ["Check", "Done", "Reading / remark"], 3,
         widths=[12.0, 3.5, 12.0], row_cm=0.62)

    signoff(doc)


def guidance(doc, logo):
    """
    What came off the forms.

    A form filled in by somebody in gloves at a machine wants columns and lines
    and nothing else; the reasoning behind the columns is worth writing down and
    is worth keeping off the form. So it is one page at the back of the pack -
    pinned up once, rather than reprinted with every shift.
    """
    masthead(doc, "HOW TO FILL THESE IN", "One page, for the notice board - the sheets "
                                          "themselves are kept clear on purpose", logo)

    items = [
        ("Every sheet",
         "One sheet per line per shift. Fill it in as the shift runs, not from memory at the "
         "end. Whoever enters it into the tablet signs and dates the bottom, so a sheet and the "
         "app can always be put side by side."),
        ("Special line",
         "One row per pass. A grade goes through two to four passes and only the finishing pass "
         "is weighed - leave the weight blank on the others, or the shift's output is counted "
         "twice. Take the grades in the order they came off: Special, SuperFine, Fine, Medium, "
         "DRC, Special DRC."),
        ("Mixing two batches",
         "Put the batch the pass is filed under in Batch No. and the other one in Mixed with. "
         "Never write both numbers in one box separated by a comma: no report can read that, and "
         "it has already had to be corrected once on five passes."),
        ("Coarse line",
         "PR1 pre-refines and weighs nothing; Refiner 2 finishes and is the only machine on the "
         "line that weighs. Log PR1's hours even when the weighing happened on the next shift - "
         "the labour belongs with the weight, and a shift showing R2 alone reads at three times "
         "what the line really does."),
        ("Grinding line",
         "Truck tyre is 30# and bike tyre is 20#. The Soorya Grinder has no electricity meter "
         "and no hour meter - leave those four columns blank rather than writing a guess."),
        ("Cracker",
         "No output column: what the cracker breaks is weighed downstream at the grinders, and a "
         "weight written here is counted twice on the plant's total. The picking gang is "
         "approximate by design - \"four of them, about three hours\" is the answer wanted."),
        ("Autoclaves",
         "One row per charge, not per shift. The three clock times are the point of the sheet: "
         "pressure reached splits the heat-up off the cook, and door opened to discharged is the "
         "vessel standing open, which is dead time on a machine that only earns while it is shut "
         "and hot. Tick Paired when both vessels were charged together and one crew attended "
         "both. About 550 kg of firewood a load is what the plant reckons on; well off that is "
         "worth a line in the remarks while somebody still remembers why."),
        ("Bearings and bushes",
         "Every 2 hours on the grinding line, every 3 hours on the refiners and the coarse line. "
         "Four positions on each machine. Anything over 80 degrees is flagged and what was done "
         "about it goes in the last column. PR1, Refiner 1, Refiner 2 and Grinder 1 run on "
         "bushes; everything else on bearings."),
        ("Boiler",
         "This is the one sheet with nothing behind it in the app - the boiler is not in MANNA "
         "RECLAIM at all. Correct the columns to match how it is actually attended and they can "
         "be built into the app afterwards."),
    ]

    t = doc.add_table(rows=len(items), cols=2)
    t.autofit = False
    borders(t)
    widths = fit([5.0, 22.7])
    for i, w in enumerate(widths):
        t.columns[i].width = Cm(w)
    for row, (head, text) in zip(t.rows, items):
        row.cells[0].width = Cm(widths[0])
        row.cells[1].width = Cm(widths[1])
        shade(row.cells[0], HEAD_FILL)
        set_text(row.cells[0], head, size=8.5, bold=True, colour=ORANGE)
        set_text(row.cells[1], text, size=8)



SHEETS = (
    ("special", special_line),
    ("coarse", coarse_line),
    ("grinders", grinders),
    ("cracker", cracker),
    ("autoclaves", autoclaves),
    ("boiler", boiler),
    ("guidance", guidance),
)


def build(out_path, logo):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(9)

    landscape(doc, first=True)
    print("Height used per sheet:")
    for i, (name, draw) in enumerate(SHEETS):
        start_sheet()
        if i:
            # Each sheet is its own section, so one page can be printed alone -
            # which is what happens in practice: the grinding line does not want
            # the boiler's sheet.
            landscape(doc)
        draw(doc, logo)
        check_sheet(name)

    doc.save(out_path)
    return out_path


def main():
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--logo", default=str(Path(__file__).parent / "manna-group.png"),
                    help="PNG or JPG of the Manna Group logo. A placeholder box is drawn if "
                         "the file is not there.")
    ap.add_argument("--out", default=str(Path(__file__).parent / "Manna-Reclaim-Log-Sheets.docx"))
    args = ap.parse_args()

    logo = Path(args.logo)
    if not logo.exists():
        print(f"No logo at {logo} - drawing a placeholder box instead.")
    out = build(args.out, logo if logo.exists() else None)
    print(f"Written: {out}")


if __name__ == "__main__":
    main()
